import streamlit as st
import edge_tts
from pydub import AudioSegment
import asyncio
import io
import re
import os
import gc
import traceback
import docx
from streamlit_mic_recorder import speech_to_text

# ==========================================
# 1. పేజీ సెట్టింగ్స్ & పర్మనెంట్ స్టేట్స్
# ==========================================
st.set_page_config(
    page_title="ఆధ్యాత్మిక వాయిస్ యంత్రం", 
    layout="wide", 
    page_icon="🕉️"
)

st.header("🔱 ఆధ్యాత్మిక వాయిస్ సిస్టమ్")
st.caption("వాయిస్ కన్వర్షన్, ఫైల్ సపోర్ట్, లైవ్ మైక్రోఫోన్, PDF, Word & Copy - 100% స్టేబుల్ వెర్షన్")

# సెషన్ స్టేట్స్
if "main_text" not in st.session_state:
    st.session_state.main_text = ""
if "audio_bytes_data" not in st.session_state:
    st.session_state.audio_bytes_data = None
if "last_mic_text" not in st.session_state:
    st.session_state.last_mic_text = ""


# ==========================================
# 2. కోర్ హెల్పర్ ఫంక్షన్లు (Ultra-Reliable)
# ==========================================

async def generate_voice_file(text, voice, pitch_val, rate_val, output_filename):
    communicate = edge_tts.Communicate(text, voice, pitch=pitch_val, rate=rate_val)
    await communicate.save(output_filename)


def split_text_into_chunks(text, max_chars=300):
    clean_text = re.sub(r'\s+', ' ', text).strip()
    sentences = re.split(r'(?<=[.!?\n।])\s+', clean_text)
    chunks = []
    current_chunk = ""
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= max_chars:
            current_chunk += sentence + " "
        else:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            current_chunk = sentence + " "
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    return [c.strip() for c in chunks if len(c.strip()) > 0]


def extract_text_from_file(uploaded_file):
    extracted = ""
    if uploaded_file.name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        extracted = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif uploaded_file.name.endswith(".txt"):
        extracted = uploaded_file.read().decode("utf-8")
    return extracted


def create_docx_bytes(text):
    doc = docx.Document()
    for paragraph in text.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


# ==========================================
# 3. ఇన్‌పుట్ విభాగం (ఫైల్ అప్‌లోడ్ & మైక్రోఫోన్)
# ==========================================
st.divider()
col_file, col_mic = st.columns([0.5, 0.5])

uploaded_file = None
with col_file:
    st.markdown("**📁 మీ ఫైల్‌ను అప్‌లోడ్ చేయండి (.docx, .txt):**")
    uploaded_file = st.file_uploader(
        "గరిష్ఠ సైజు 10MB వరకు అనుకూలం", 
        type=["docx", "txt"],
        help="కేవలం Microsoft Word (.docx) లేదా Text (.txt) ఫైల్స్ మాత్రమే సపోర్ట్ చేయబడతాయి."
    )
    
    if uploaded_file is not None:
        max_mb = 10
        if uploaded_file.size > max_mb * 1024 * 1024:
            st.error(f"⚠️ ఫైల్ సైజు {max_mb} MB కంటే తక్కువగా ఉండాలి!")
        else:
            try:
                f_text = extract_text_from_file(uploaded_file)
                if f_text and f_text != st.session_state.main_text:
                    st.session_state.main_text = f_text
                    st.success(f"✅ '{uploaded_file.name}' విజయవంతంగా లోడ్ అయింది!")
            except Exception as fe:
                st.error(f"ఫైల్ చదవడంలో లోపం: {fe}")

with col_mic:
    st.markdown("**🎙️ మైక్రోఫోన్ ద్వారా మాట్లాడండి (Live Voice Typing):**")
    mic_lang = st.selectbox("మాట్లాడే భాష:", options=["తెలుగు (Telugu)", "హిందీ (Hindi)", "ఇంగ్లీష్ (English)"])
    mic_code_map = {"తెలుగు (Telugu)": "te-IN", "హిందీ (Hindi)": "hi-IN", "ఇంగ్లీష్ (English)": "en-IN"}
    
    spoken_result = speech_to_text(
        start_prompt="🎙️ మాట్లాడటం ప్రారంభించండి (Start)",
        stop_prompt="⏹️ ఆపండి (Stop)",
        language=mic_code_map[mic_lang],
        use_container_width=True,
        key='perfect_mic_recorder'
    )
    
    if spoken_result and spoken_result != st.session_state.last_mic_text:
        st.session_state.main_text += " " + spoken_result
        st.session_state.last_mic_text = spoken_result
        st.rerun()

# ప్రధాన టెక్స్ట్ ఏరియా
user_input_text = st.text_area(
    "ఆడియో/ఫైల్స్‌గా మార్చాలనుకుంటున్న టెక్స్ట్:", 
    value=st.session_state.main_text, 
    height=180,
    placeholder="ఇక్కడ టెక్స్ట్ పేస్ట్ చేయండి లేదా పైన ఉన్న మైక్రోఫోన్ / ఫైల్ అప్‌లోడ్ ఉపయోగించండి..."
)

if user_input_text != st.session_state.main_text:
    st.session_state.main_text = user_input_text


# ==========================================
# 4. ఆడియో ఎంపికలు & ఆప్షనల్ కంట్రోల్స్
# ==========================================
st.divider()
col_lang, col_voice = st.columns([0.5, 0.5])

with col_lang:
    selected_lang = st.selectbox("🌐 ఆడియో భాషను ఎంచుకోండి:", options=["తెలుగు (Telugu)", "హిందీ (Hindi)", "ఇంగ్లీష్ (English)"])

with col_voice:
    if "తెలుగు" in selected_lang:
        voice_option = st.radio("🎙️ స్వరాన్ని ఎంచుకోండి:", options=["👨 మోహన్ (పురుష)", "👩 శ్రుతి (స్త్రీ)"], horizontal=True)
    elif "హిందీ" in selected_lang:
        voice_option = st.radio("🎙️ స్వరాన్ని ఎంచుకోండి:", options=["👨 మధుర్ (పురుష)", "👩 స్వర్ణ (స్త్రీ)"], horizontal=True)
    else:
        voice_option = st.radio("🎙️ స్వరాన్ని ఎంచుకోండి:", options=["👨 ప్రభాత్ (పురుష)", "👩 నీరజ (స్త్రీ)"], horizontal=True)

# 🎛️ ఆప్షనల్ ఆడియో సెట్టింగ్స్
with st.expander("⚙️ ఆప్షనల్ ఆడియో సెట్టింగ్స్ (స్పీడ్, పిచ్ & BGM ఫైన్-ట్యూనింగ్)"):
    col_opt_speed, col_opt_pitch, col_opt_pause = st.columns(3)
    
    with col_opt_speed:
        audio_speed = st.select_slider("🔊 ప్లే స్పీడ్ (Play Speed):", options=[0.75, 0.85, 1.0, 1.15, 1.25, 1.5], value=0.85)
    
    with col_opt_pitch:
        pitch_custom = st.select_slider("🎚️ వాయిస్ గంభీరత (Pitch/Base):", options=["సాధారణ (Normal)", "గంభీరం (Deep Base)", "అత్యంత గంభీరం (Heavy Base)"], value="సాధారణ (Normal)")
        
    with col_opt_pause:
        pause_duration = st.slider("⏸️ వాక్యాల మధ్య విరామం (Pause Sec):", min_value=0.3, max_value=2.0, value=0.6, step=0.1)
        
    col_bgm_1, col_bgm_2 = st.columns([0.4, 0.6])
    with col_bgm_1:
        enable_bgm = st.checkbox("🎶 BGM (బ్యాక్‌గ్రౌండ్ మ్యూజిక్) జోడించు", value=True)
    with col_bgm_2:
        bgm_volume = st.slider("🎵 BGM శబ్దం (Volume %):", min_value=2, max_value=20, value=6)


# ==========================================
# 5. ఐదు ప్రధాన ఆప్షన్ల వరుస (Action Controls)
# ==========================================
st.markdown("##### 🎯 యాక్షన్ కంట్రోల్స్ (Action Controls)")

active_text = st.session_state.main_text.strip()

col_btn1, col_btn2, col_btn3, col_btn4, col_btn5 = st.columns([0.22, 0.18, 0.18, 0.22, 0.20])

with col_btn1:
    convert_btn = st.button("🔊 ఆడియో చేయి", type="primary", use_container_width=True)

with col_btn2:
    if active_text:
        pdf_html = f"<html><head><meta charset='utf-8'></head><body><p style='font-size:16px;'>{active_text.replace('\n', '<br>')}</p></body></html>"
        st.download_button(
            label="📄 PDF ఫైల్",
            data=pdf_html.encode('utf-8'),
            file_name="spiritual_note.html",
            mime="text/html",
            use_container_width=True
        )
    else:
        st.button("📄 PDF ఫైల్", disabled=True, use_container_width=True)

with col_btn3:
    if active_text:
        docx_data = create_docx_bytes(active_text)
        st.download_button(
            label="📝 Word ఫైల్",
            data=docx_data,
            file_name="spiritual_note.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.button("📝 Word ఫైల్", disabled=True, use_container_width=True)

with col_btn4:
    if active_text:
        if st.button("📋 టెక్స్ట్ కాపీ", use_container_width=True):
            st.code(active_text, language=None)
            st.toast("✅ పైన ఉన్న టెక్స్ట్‌ని క్లిక్ చేసి కాపీ చేసుకోండి!", icon="📋")
    else:
        st.button("📋 టెక్స్ట్ కాపీ", disabled=True, use_container_width=True)

with col_btn5:
    if st.button("🧹 మొత్తం క్లియర్", use_container_width=True):
        st.session_state.main_text = ""
        st.session_state.audio_bytes_data = None
        st.session_state.last_mic_text = ""
        gc.collect()
        st.rerun()


# ==========================================
# 6. హై-స్పీడ్ ఆడియో ప్రాసెసింగ్ లాజిక్
# ==========================================
if convert_btn:
    if active_text:
        with st.spinner("ఆడియో వేగంగా ప్రాసెస్ అవుతోంది... దయచేసి వేచి ఉండండి..."):
            try:
                clean_txt = re.sub(r'[*#_~`]', '', active_text)
                
                voice_map = {
                    "👨 మోహన్ (పురుష)": "te-IN-MohanNeural",
                    "👩 శ్రుతి (స్త్రీ)": "te-IN-ShrutiNeural",
                    "👨 మధుర్ (పురుష)": "hi-IN-MadhurNeural",
                    "👩 స్వర్ణ (స్త్రీ)": "hi-IN-SwaraNeural",
                    "👨 ప్రభాత్ (పురుష)": "en-IN-PrabhatNeural",
                    "👩 నీరజ (స్త్రీ)": "en-IN-NeerjaNeural"
                }
                selected_voice = voice_map[voice_option]

                rate_str = f"{int((audio_speed - 1.0) * 100):+d}%"
                pitch_val_map = {
                    "సాధారణ (Normal)": "+0Hz",
                    "గంభీరం (Deep Base)": "-5Hz",
                    "అత్యంత గంభీరం (Heavy Base)": "-10Hz"
                }
                pitch_str = pitch_val_map[pitch_custom]

                text_chunks = split_text_into_chunks(clean_txt, max_chars=300)
                speech_sound = AudioSegment.empty()
                silence_pause = AudioSegment.silent(duration=int(pause_duration * 1000))

                for i, chunk in enumerate(text_chunks):
                    temp_file = f"temp_{i}.mp3"
                    try:
                        asyncio.run(generate_voice_file(chunk, selected_voice, pitch_str, rate_str, temp_file))
                        if os.path.exists(temp_file) and os.path.getsize(temp_file) > 0:
                            chunk_sound = AudioSegment.from_file(temp_file)
                            speech_sound += chunk_sound + silence_pause
                            os.remove(temp_file)
                    except Exception:
                        pass

                if len(speech_sound) > 0:
                    final_sound = speech_sound
                    if enable_bgm and os.path.exists("bgm.mp3"):
                        try:
                            bgm_sound = AudioSegment.from_file("bgm.mp3")
                            if len(bgm_sound) < len(speech_sound):
                                bgm_sound = bgm_sound * ((len(speech_sound) // len(bgm_sound)) + 1)
                            
                            bgm_sound = bgm_sound[:len(speech_sound) + 1000]
                            reduction_db = 22 - (bgm_volume * 1.5)
                            bgm_sound = bgm_sound - reduction_db
                            final_sound = speech_sound.overlay(bgm_sound)
                        except Exception:
                            pass

                    final_fp = io.BytesIO()
                    final_sound.export(final_fp, format="mp3")
                    st.session_state.audio_bytes_data = final_fp.getvalue()
                    gc.collect()
                    st.success("🎉 ఆడియో విజయవంతంగా సిద్ధమైంది!")
                else:
                    st.error("❌ ఆడియో డేటా ఏదీ జనరేట్ కాలేదు!")

            except Exception as e:
                st.error("❌ ఆడియో సిస్టమ్‌లో లోపం వచ్చింది:")
                st.code(traceback.format_exc())
    else:
        st.warning("దయచేసి టెక్స్ట్ ఎంటర్ చేయండి లేదా మాట్లాడండి.")

# 📥 స్థిరమైన ఆడియో ప్లేయర్
if st.session_state.audio_bytes_data is not None:
    st.divider()
    st.audio(st.session_state.audio_bytes_data, format="audio/mp3")
    st.download_button(
        label="📥 MP3 ఆడియో ఫైల్‌ని డౌన్‌లోడ్ చేయండి", 
        data=st.session_state.audio_bytes_data, 
        file_name="spiritual_audio.mp3", 
        mime="audio/mp3",
        key="permanent_download_btn",
        use_container_width=True
    )
